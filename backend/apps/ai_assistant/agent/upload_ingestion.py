"""
文件 → 原始文档 → 分块 入库管线

支持的文件格式：
  文本类：.txt / .md / .csv / .json / .xml / .yaml / .yml / .html / .htm
  文档类：.pdf / .docx
  压缩包：.zip / .tar.gz / .tgz / .tar
         （RAR 不在标准库内，跳过并报错提示）

用法：
    from .upload_ingestion import ingest_uploaded_file
    docs = ingest_uploaded_file(request.FILES["file"], user=request.user)
"""

from __future__ import annotations

import hashlib
import io
import logging
import os
import re
import tarfile
import tempfile
import zipfile
from pathlib import Path
from typing import Iterable, List, Optional, Tuple

logger = logging.getLogger("apps.ai_assistant")

SUPPORTED_EXTS = {
    ".txt", ".md", ".csv", ".json", ".xml", ".yaml", ".yml", ".html", ".htm",
    ".pdf", ".docx",
}
ARCHIVE_EXTS = {".zip", ".tar.gz", ".tgz", ".tar"}


class UploadIngestionError(Exception):
    """用户可直接展示的导入错误"""


# ---------------------------------------------------------------------------
# 解析器：字节流 → 纯文本 + 元数据
# ---------------------------------------------------------------------------

def _parse_plain_bytes(raw: bytes, source_type: str) -> Tuple[str, dict]:
    encodings = ["utf-8", "gbk", "gb18030", "utf-16", "latin-1"]
    last_err = None
    for enc in encodings:
        try:
            text = raw.decode(enc)
            return text, {"encoding": enc, "bytes": len(raw)}
        except UnicodeDecodeError as exc:
            last_err = exc
    raise UploadIngestionError(f"无法解码 {source_type} 文件：{last_err}")


def _parse_pdf(raw: bytes, _ext: str = "") -> Tuple[str, dict]:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise UploadIngestionError("缺少 PyPDF2 依赖，请 pip install PyPDF2")
    reader = PdfReader(io.BytesIO(raw))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:
            pages.append("")
            logger.warning("[pdf] page extract failed: %s", exc)
    text = "\n\n".join(pages)
    metadata = {"pages": len(reader.pages), "parser": "PyPDF2"}
    return text, metadata


def _parse_docx(raw: bytes, _ext: str = "") -> Tuple[str, dict]:
    try:
        from docx import Document
    except ImportError:
        raise UploadIngestionError("缺少 python-docx 依赖，请 pip install python-docx")
    doc = Document(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return text, {"parser": "python-docx"}


def _parse_html(raw: bytes, _ext: str = "") -> Tuple[str, dict]:
    import html
    try:
        import lxml.html  # 可选
        node = lxml.html.fromstring(raw.decode("utf-8", errors="ignore"))
        text = node.xpath("string()") if False else node.text_content()
    except ImportError:
        text = re.sub(r"<script[^>]*>.*?</script>", " ", raw.decode("utf-8", errors="ignore"), flags=re.S)
        text = re.sub(r"<style[^>]*>.*?</style>", " ", text, flags=re.S)
        text = re.sub(r"<[^>]+>", " ", text)
        text = html.unescape(text)
    text = re.sub(r"\s+", " ", text).strip()
    return text, {"parser": "html"}


PARSERS = {
    ".txt": _parse_plain_bytes,
    ".md": _parse_plain_bytes,
    ".csv": _parse_plain_bytes,
    ".json": _parse_plain_bytes,
    ".xml": _parse_plain_bytes,
    ".yaml": _parse_plain_bytes,
    ".yml": _parse_plain_bytes,
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".html": _parse_html,
    ".htm": _parse_html,
}


def detect_source_type(name: str) -> str:
    lower = name.lower()
    for ext in ARCHIVE_EXTS:
        if lower.endswith(ext):
            return "archive"
    for ext in PARSERS:
        if lower.endswith(ext):
            return ext.lstrip(".")
    return "unknown"


def _parse_one(name: str, raw: bytes) -> Tuple[str, dict]:
    lower = name.lower()
    parser = None
    ext = None
    for e, fn in PARSERS.items():
        if lower.endswith(e):
            parser = fn
            ext = e
            break
    if parser is None:
        # 不支持的格式 → 只按 UTF-8 读入，报错由调用方决定
        raise UploadIngestionError(f"不支持的文件类型：{name.lower().rsplit('.', 1)[-1]}")
    text, meta = parser(raw, ext)
    return text, meta


# ---------------------------------------------------------------------------
# 压缩包解包：返回 [(inner_name, raw_bytes), ...]
# ---------------------------------------------------------------------------

def _unpack_archive(raw: bytes, archive_name: str) -> List[Tuple[str, bytes]]:
    """安全解包，跳过 __MACOSX 等垃圾目录、零字节和不可解析文件。"""
    lower = archive_name.lower()
    results: List[Tuple[str, bytes]] = []
    seen = set()

    def _is_allowed(name: str) -> bool:
        base = os.path.basename(name)
        if base.startswith(".") or base.startswith("~$"):
            return False
        if "/__MACOSX/" in name or "\\__MACOSX\\" in name:
            return False
        ext = Path(name).suffix.lower()
        return ext in SUPPORTED_EXTS

    with tempfile.TemporaryDirectory(prefix="rag_unpack_") as tmpdir:
        tmp = Path(tmpdir)
        archive_path = tmp / archive_name
        archive_path.write_bytes(raw)

        try:
            if lower.endswith(".zip"):
                with zipfile.ZipFile(archive_path) as zf:
                    zf.extractall(tmp)
            elif lower.endswith(".tar.gz") or lower.endswith(".tgz"):
                with tarfile.open(archive_path, "r:gz") as tf:
                    tf.extractall(tmp)
            elif lower.endswith(".tar"):
                with tarfile.open(archive_path, "r:") as tf:
                    tf.extractall(tmp)
            elif lower.endswith(".rar"):
                raise UploadIngestionError("RAR 压缩包暂不支持，请改用 ZIP 或 TAR.GZ")
            else:
                raise UploadIngestionError(f"不支持的压缩包格式：{lower}")
        except UploadIngestionError:
            raise
        except Exception as exc:
            raise UploadIngestionError(f"压缩包解包失败：{exc}") from exc

        for p in tmp.rglob("*"):
            if not p.is_file():
                continue
            rel = str(p.relative_to(tmp))
            if not _is_allowed(rel):
                continue
            if rel in seen:
                continue
            seen.add(rel)
            try:
                results.append((rel, p.read_bytes()))
            except Exception:
                continue

    return results


# ---------------------------------------------------------------------------
# 文本分块 + 关键词抽取
# ---------------------------------------------------------------------------

def _split_into_chunks(text: str, target_chars: int = 800, overlap: int = 120) -> List[str]:
    """按段落边界优先切分，不足则软切。"""
    if not text or not text.strip():
        return []
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: List[str] = []
    buf = ""
    for para in paragraphs:
        if len(buf) + len(para) + 2 <= target_chars:
            buf = f"{buf}\n\n{para}" if buf else para
        else:
            if buf:
                chunks.append(buf)
            if len(para) > target_chars:
                # 长段落按句子切
                sentences = re.split(r"(?<=[。！？!?])\s*", para)
                buf = ""
                for sent in sentences:
                    if len(buf) + len(sent) <= target_chars:
                        buf = f"{buf}{sent}" if buf else sent
                    else:
                        if buf:
                            chunks.append(buf)
                        if len(sent) > target_chars:
                            for i in range(0, len(sent), target_chars):
                                chunks.append(sent[i:i + target_chars])
                            buf = ""
                        else:
                            buf = sent
            else:
                buf = para
    if buf:
        chunks.append(buf)

    # 去重叠后返回
    final = []
    prev = ""
    for c in chunks:
        if prev and overlap > 0 and len(c) > overlap:
            c = prev[-overlap:] + c
        final.append(c.strip())
        prev = c
    return final


def _extract_keywords(text: str, top_n: int = 8) -> List[str]:
    try:
        import jieba
        from collections import Counter
        stopwords = {"的", "了", "是", "在", "和", "与", "或", "不", "也", "为", "对", "及", "等", "以", "将", "应", "可", "有", "没", "就", "都", "而", "但", "如果", "因此", "所以", "一个", "一些", "这种", "那种"}
        words = [w.strip() for w in jieba.cut(text) if len(w.strip()) > 1 and w.strip() not in stopwords]
        return [w for w, _ in Counter(words).most_common(top_n)]
    except ImportError:
        return []


def _estimate_tokens(text: str) -> int:
    return max(1, len(text) // 3)


def _make_summary(text: str, max_len: int = 160) -> str:
    """取前几段中非空内容作为摘要"""
    first_lines = []
    for line in text.splitlines():
        s = line.strip()
        if s:
            first_lines.append(s)
            if sum(len(l) for l in first_lines) > max_len * 2:
                break
    summary = " ".join(first_lines)
    return summary[:max_len] + ("..." if len(summary) > max_len else "")


# ---------------------------------------------------------------------------
# 主入口
# ---------------------------------------------------------------------------

def _sha256(raw: bytes) -> str:
    return hashlib.sha256(raw).hexdigest()


def _ingest_text_into_chunks(
    doc,
    text: str,
    chunk_size: int = 800,
    overlap: int = 120,
) -> int:
    """把解析后的文本切成 AiKnowledgeChunk 并关联到 doc。返回分块数。"""
    from apps.ai_assistant.models import AiKnowledgeChunk

    if doc.status == "ready" and doc.chunks.exists():
        logger.info("[ingest] doc %d already has chunks, skipping", doc.id)
        return doc.chunks.count()

    if not text or not text.strip():
        doc.status = "failed"
        doc.error_message = "文档解析后内容为空"
        doc.save(update_fields=["status", "error_message"])
        return 0

    chunks = _split_into_chunks(text, target_chars=chunk_size, overlap=overlap)
    # 清理旧分块（幂等）
    doc.chunks.all().delete()

    objs = []
    for idx, chunk_text in enumerate(chunks, start=1):
        objs.append(AiKnowledgeChunk(
            document=doc,
            content=chunk_text,
            ordinal=idx,
            token_estimate=_estimate_tokens(chunk_text),
            keywords=_extract_keywords(chunk_text),
        ))
    AiKnowledgeChunk.objects.bulk_create(objs)

    doc.status = "ready"
    doc.error_message = ""
    doc.metadata = {
        **(doc.metadata or {}),
        "chunk_count": len(objs),
        "total_chars": len(text),
        "summary": _make_summary(text),
    }
    doc.save(update_fields=["status", "error_message", "metadata"])
    return len(objs)


def _sanitize_text(text: str) -> str:
    """PostgreSQL 不允许 NUL (0x00)，同时清理控制字符和过长空白。"""
    if not text:
        return text
    text = text.replace("\x00", "")
    text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _create_document_record(
    name: str,
    raw: bytes,
    *,
    user=None,
    region=None,
    is_public: bool = True,
    title: str = "",
) -> "AiKnowledgeDocument":
    from apps.ai_assistant.models import AiKnowledgeDocument

    checksum = _sha256(raw)
    source_type = detect_source_type(name)
    ext = re.sub(r"^archive$", "", source_type) or "bin"

    existing = AiKnowledgeDocument.objects.filter(checksum=checksum).first()
    if existing and existing.status == "ready":
        logger.info("[ingest] duplicate skipped: %s (checksum=%s)", name, checksum[:12])
        return existing

    title_clean = (title or "").strip() or Path(name).stem
    title_clean = re.sub(r"^[【\(（].*?[】\)）]\s*", "", title_clean)
    title_clean = title_clean[:200]

    doc = existing or AiKnowledgeDocument(
        title=title_clean,
        source_type=ext,
        checksum=checksum,
        parser="text",
        region=region,
        is_public=is_public,
        created_by=user,
        status="pending",
    )
    doc.source_path = name

    # 如果是 Django UploadedFile 实例带 .file，让它保存文件
    doc.file.save(Path(name).name, io.BytesIO(raw), save=False)

    try:
        text, meta = _parse_one(name, raw)
    except UploadIngestionError as exc:
        doc.status = "failed"
        doc.error_message = str(exc)
        doc.parser = "none"
        doc.save()
        return doc

    text = _sanitize_text(text)

    doc.parser = meta.get("parser", source_type)
    doc.parser_version = meta.get("version", "")
    doc.metadata = meta

    doc.save()
    _ingest_text_into_chunks(doc, text)
    return doc


def ingest_uploaded_file(
    upload_file,
    *,
    user=None,
    title: str = "",
    region=None,
    is_public: bool = True,
) -> List["AiKnowledgeDocument"]:
    """
    上传入口：
      单文件（txt/md/pdf/docx/html/json/yaml/xml）→ 直接入库
      压缩包（zip / tar.gz / tgz / tar）→ 解包后逐文件入库
    返回创建/复用的 AiKnowledgeDocument 列表。
    """
    name = getattr(upload_file, "name", str(upload_file))
    raw = upload_file.read()
    upload_file.seek(0)

    source_type = detect_source_type(name)

    if source_type == "archive":
        inner_files = _unpack_archive(raw, name)
        if not inner_files:
            raise UploadIngestionError(f"压缩包 {name} 中没有可解析的文件")
        docs: List["AiKnowledgeDocument"] = []
        errors: List[str] = []
        for inner_name, inner_raw in inner_files:
            try:
                d = _create_document_record(
                    inner_name, inner_raw,
                    user=user, region=region, is_public=is_public,
                )
                docs.append(d)
            except UploadIngestionError as exc:
                errors.append(f"{inner_name}: {exc}")
                logger.warning("[ingest] inner failed: %s -> %s", inner_name, exc)
        if not docs:
            raise UploadIngestionError(
                f"压缩包内文件全部解析失败：{'; '.join(errors[:3])}"
            )
        if errors:
            logger.warning("[ingest] archive partial failures: %d/%d",
                           len(errors), len(inner_files))
        return docs

    # 单文件
    doc = _create_document_record(
        name, raw,
        user=user, region=region, is_public=is_public, title=title,
    )
    return [doc]
